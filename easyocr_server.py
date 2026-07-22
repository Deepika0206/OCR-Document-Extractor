import base64
import io
import json
import re
import threading
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

import docx
import easyocr
import fitz  # PyMuPDF

from reading_order import reconstruct_reading_order, detect_redaction_boxes

PORT = 5001
MIN_TEXT_LAYER_CHARS = 20
MAX_PDF_PAGES = 25
print("Loading EasyOCR model (first run downloads ~65MB of model weights, then it's cached)...", flush=True)
reader = easyocr.Reader(["en"], gpu=False)
print("EasyOCR ready.", flush=True)
inference_lock = threading.Lock()

DATA_URI_RE = re.compile(r"^data:([a-zA-Z0-9.+/-]+);base64,(.+)$", re.DOTALL)
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

def ocr_image_bytes(image_bytes: bytes):
    with inference_lock:
        results = reader.readtext(image_bytes)
    try:
        results = results + detect_redaction_boxes(image_bytes, results)
    except Exception:
    
        pass
    return reconstruct_reading_order(results)


def extract_pdf(pdf_bytes: bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        total_pages = doc.page_count
        pages_to_process = min(total_pages, MAX_PDF_PAGES)
        parts = []
        pages_meta = []
        for page_index in range(pages_to_process):
            page = doc[page_index]
            page_text = page.get_text().strip()
            if len(page_text) >= MIN_TEXT_LAYER_CHARS:
                parts.append(page_text)
                pages_meta.append({"page": page_index + 1, "source": "text_layer"})
            else:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                page_image_bytes = pix.tobytes("png")
                page_text, page_ocr_meta = ocr_image_bytes(page_image_bytes)
                parts.append(page_text)
                pages_meta.append({"page": page_index + 1, "source": "ocr", **page_ocr_meta})
        text = "\n\n".join(p for p in parts if p)
        if total_pages > MAX_PDF_PAGES:
            text += f"\n\n[Note: this PDF has {total_pages} pages; only the first {MAX_PDF_PAGES} were processed.]"
        return text, {"pages": pages_meta}
    finally:
        doc.close()


def extract_docx(docx_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(docx_bytes))
    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


class OCRHandler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        pass

    def _send_json(self, status, payload):
        body = json.dumps(payload).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_POST(self):
        if self.path == "/build-docx":
            self._handle_build_docx()
            return
        if self.path != "/extract":
            self._send_json(404, {"error": "Not found"})
            return

        try:
            content_length = int(self.headers.get("Content-Length", 0))
            raw_body = self.rfile.read(content_length)
            data = json.loads(raw_body.decode("utf-8")) if raw_body else {}
        except Exception:
            self._send_json(400, {"error": "Invalid request body"})
            return

        file_data_uri = data.get("file", "")
        match = DATA_URI_RE.match(file_data_uri or "")
        if not match:
            self._send_json(400, {"error": "Invalid file format. Expected base64 Data URI."})
            return

        mime_type = match.group(1)
        try:
            file_bytes = base64.b64decode(match.group(2))
        except Exception:
            self._send_json(400, {"error": "Could not decode base64 file data."})
            return

        meta = None
        try:
            if mime_type.startswith("image/"):
                text, meta = ocr_image_bytes(file_bytes)
            elif mime_type == "application/pdf":
                text, meta = extract_pdf(file_bytes)
            elif mime_type == DOCX_MIME:
                text = extract_docx(file_bytes)
            elif mime_type == "text/plain":
                text = file_bytes.decode("utf-8", errors="replace")
            else:
                self._send_json(400, {
                    "error": f"Unsupported file type: {mime_type}. "
                             "Supported: images (PNG/JPG/WebP), PDF, and Word (.docx) documents."
                })
                return
        except Exception as e:
            self._send_json(500, {"error": f"Extraction failed: {e}"})
            return

        response = {"text": text}
        if meta is not None:
            response["meta"] = meta
        self._send_json(200, response)

    def _handle_build_docx(self):
        try:
            content_length = int(self.headers.get("Content-Length", 0))
            raw_body = self.rfile.read(content_length)
            data = json.loads(raw_body.decode("utf-8")) if raw_body else {}
        except Exception:
            self._send_json(400, {"error": "Invalid request body"})
            return

        text = data.get("text", "")
        if not isinstance(text, str) or not text.strip():
            self._send_json(400, {"error": "No text provided to export."})
            return

        try:
            document = docx.Document()
            for para in text.split("\n\n"):
                para = para.strip("\n")
                if not para.strip():
                    continue
                # Preserve single line breaks within a paragraph as line
                # breaks in the Word doc, rather than collapsing them.
                p = document.add_paragraph()
                lines = para.split("\n")
                for i, line in enumerate(lines):
                    if i > 0:
                        p.add_run().add_break()
                    p.add_run(line)
            buf = io.BytesIO()
            document.save(buf)
            docx_base64 = base64.b64encode(buf.getvalue()).decode("ascii")
        except Exception as e:
            self._send_json(500, {"error": f"Failed to build Word document: {e}"})
            return

        self._send_json(200, {"docx_base64": docx_base64})

    def do_GET(self):
        if self.path == "/health":
            self._send_json(200, {"status": "ready"})
        else:
            self._send_json(404, {"error": "Not found"})


if __name__ == "__main__":
    server = ThreadingHTTPServer(("127.0.0.1", PORT), OCRHandler)
    print(f"Local extraction backend listening on http://127.0.0.1:{PORT}", flush=True)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nShutting down.", flush=True)